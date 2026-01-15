from docx import Document


template_path = 'template/template.docx'
data_keys = ['full_name', 'birth_year', 'gender', 'education', 'edu_start',
             'edu_end', "study_format", 'has_experience', 'position', 'company', 'work_period',
             "currently_working", 'uzbek', 'russian', 'other_langs', 'family', 
             'phone', 'telegram']


def write_template(data):
    doc = Document(template_path)
    for i in range(2,20):
        # doc.paragraphs[i].runs[1].text = data[data_keys[i-2]]
        doc.paragraphs[i].text = doc.paragraphs[i].text.replace('$', data[data_keys[i-2]])

    
    doc.save(f"template/{data["full_name"]}.docx")